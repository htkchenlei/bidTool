import os
import re
import tempfile
import subprocess
from typing import List, Dict, Tuple, Optional


def convert_to_markdown(file_path: str, file_type: str) -> Dict:
    """
    使用 markitdown 将文档转换为 markdown 格式

    :param file_path: 文件路径
    :param file_type: 文件类型 (DOCX/PDF/IMAGE/TEXT)
    :return: {
        "markdown": "转换后的markdown文本",
        "sections": [{"level": int, "title": str, "content": str}],
        "success": bool,
        "error": str or None
    }
    """
    result = {
        "markdown": "",
        "sections": [],
        "success": False,
        "error": None
    }

    if not os.path.exists(file_path):
        result["error"] = f"文件不存在: {file_path}"
        return result

    try:
        if file_type == "DOCX":
            markdown = _convert_docx_to_markdown(file_path)
        elif file_type == "PDF":
            markdown = _convert_pdf_to_markdown(file_path)
        elif file_type == "IMAGE":
            markdown = _convert_image_to_markdown(file_path)
        elif file_type == "TEXT":
            markdown = _convert_text_to_markdown(file_path)
        else:
            result["error"] = f"不支持的文件类型: {file_type}"
            return result

        if not markdown.strip():
            result["error"] = "转换结果为空"
            return result

        result["markdown"] = markdown
        result["sections"] = extract_sections(markdown)
        result["success"] = True

    except Exception as e:
        result["error"] = f"转换失败: {str(e)}"

    return result


def _convert_docx_to_markdown(file_path: str) -> str:
    """使用 markitdown 转换 DOCX 文件"""
    try:
        from markitdown import convert
        result = convert(file_path, output_format="markdown")
        if result and hasattr(result, 'text'):
            return result.text
        elif isinstance(result, str):
            return result
        elif isinstance(result, dict) and 'markdown' in result:
            return result['markdown']
        return str(result)
    except ImportError:
        return _fallback_docx_to_markdown(file_path)
    except Exception:
        return _fallback_docx_to_markdown(file_path)


def _convert_pdf_to_markdown(file_path: str) -> str:
    """使用 markitdown 转换 PDF 文件"""
    try:
        from markitdown import convert
        result = convert(file_path, output_format="markdown")
        if result and hasattr(result, 'text'):
            return result.text
        elif isinstance(result, str):
            return result
        elif isinstance(result, dict) and 'markdown' in result:
            return result['markdown']
        return str(result)
    except ImportError:
        return _fallback_pdf_to_markdown(file_path)
    except Exception:
        return _fallback_pdf_to_markdown(file_path)


def _convert_image_to_markdown(file_path: str) -> str:
    """图片文件转换为 markdown（使用 OCR）"""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img, lang="chi_sim+eng")
        if text.strip():
            return f"## 图片内容\n\n{text.strip()}"
        return ""
    except Exception:
        return f"![图片]({os.path.basename(file_path)})"


def _convert_text_to_markdown(file_path: str) -> str:
    """纯文本文件转换为 markdown"""
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
    return content


def _fallback_docx_to_markdown(file_path: str) -> str:
    """DOCX 转换兜底方案（使用 python-docx）"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                style = para.style.name.lower()
                if 'heading' in style or '标题' in style:
                    level = 1
                    if '2' in style or '二' in style:
                        level = 2
                    elif '3' in style or '三' in style:
                        level = 3
                    elif '4' in style or '四' in style:
                        level = 4
                    paragraphs.append('#' * level + ' ' + para.text.strip())
                else:
                    paragraphs.append(para.text.strip())

        for table in doc.tables:
            table_lines = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_lines.append('| ' + ' | '.join(row_data) + ' |')
            if table_lines:
                header = table_lines[0]
                separator = '| ' + ' | '.join(['---'] * len(header.split('|'))) + ' |'
                paragraphs.append('\n'.join([header, separator] + table_lines[1:]))

        return '\n\n'.join(paragraphs)
    except Exception as e:
        return f"文档转换失败: {str(e)}"


def _fallback_pdf_to_markdown(file_path: str) -> str:
    """PDF 转换兜底方案（使用 pdfplumber）"""
    try:
        import pdfplumber
        content = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    content.append(text)

                tables = page.extract_tables()
                for table in tables:
                    table_lines = []
                    for row in table:
                        row_data = [str(cell).strip() if cell else '' for cell in row]
                        table_lines.append('| ' + ' | '.join(row_data) + ' |')
                    if table_lines:
                        header = table_lines[0]
                        col_count = len(header.split('|')) - 2
                        separator = '| ' + ' | '.join(['---'] * col_count) + ' |'
                        content.append('\n'.join([header, separator] + table_lines[1:]))

        return '\n\n'.join(content)
    except Exception as e:
        return f"PDF 转换失败: {str(e)}"


def extract_sections(markdown_text: str) -> List[Dict]:
    """
    从 markdown 文本中提取章节结构

    :param markdown_text: markdown 格式文本
    :return: [{level, title, content, start_line, end_line}]
    """
    sections = []
    lines = markdown_text.split('\n')
    current_section = None
    current_content = []

    for idx, line in enumerate(lines):
        stripped = line.strip()

        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            if current_section:
                current_section['content'] = '\n'.join(current_content).strip()
                current_section['end_line'] = idx - 1
                sections.append(current_section)

            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            current_section = {
                'level': level,
                'title': title,
                'content': '',
                'start_line': idx,
                'end_line': idx
            }
            current_content = []
        elif current_section:
            current_content.append(line)

    if current_section:
        current_section['content'] = '\n'.join(current_content).strip()
        current_section['end_line'] = len(lines) - 1
        sections.append(current_section)

    return sections


def detect_file_type(file_path: str) -> str:
    """
    根据文件扩展名检测文件类型

    :param file_path: 文件路径
    :return: 文件类型 (DOCX/PDF/IMAGE/TEXT/UNKNOWN)
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext in ['.docx']:
        return 'DOCX'
    elif ext in ['.pdf']:
        return 'PDF'
    elif ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.webp']:
        return 'IMAGE'
    elif ext in ['.txt', '.md', '.markdown']:
        return 'TEXT'
    return 'UNKNOWN'


def extract_text_from_markdown(markdown_text: str) -> str:
    """
    从 markdown 中提取纯文本（去除格式标记）

    :param markdown_text: markdown 格式文本
    :return: 纯文本
    """
    text = markdown_text

    text = re.sub(r'#{1,6}\s+', '', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


def find_score_criteria_section(sections: List[Dict]) -> Optional[Dict]:
    """
    在章节列表中查找评分标准相关章节

    :param sections: 章节列表
    :return: 评分标准章节或 None
    """
    keywords = ['评分标准', '评分办法', '评标办法', '评审标准', '分值设置', '打分标准']

    for section in sections:
        title = section.get('title', '')
        for kw in keywords:
            if kw in title:
                return section

    for section in sections:
        content = section.get('content', '')
        for kw in keywords:
            if kw in content:
                return section

    return None
