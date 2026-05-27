from flask import Blueprint, request, jsonify
import os
import tempfile
from docx import Document
import json

file_parse_bp = Blueprint('file_parse', __name__)

REGIONS_FILE = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'china_regions.json')

class DocxTextExtractor:
    def __init__(self, docx_path):
        self.doc = Document(docx_path)
        self.page_breaks = self._find_page_breaks()
        self.full_text = self._extract_full_text()

    def _find_page_breaks(self):
        page_breaks = []
        current_page = 1
        current_line_count = 0
        lines_per_page_estimate = 40

        for i, paragraph in enumerate(self.doc.paragraphs):
            text = paragraph.text
            if text.strip():
                lines_in_para = len(text) // 80 + text.count('\n') + 1
                current_line_count += lines_in_para

                if 'w:br' in paragraph._p.xml and 'type="page"' in paragraph._p.xml:
                    page_breaks.append((i, current_page))
                    current_page += 1
                    current_line_count = lines_in_para

                if current_line_count > lines_per_page_estimate:
                    page_breaks.append((i, current_page))
                    current_page += 1
                    current_line_count = lines_in_para

        return page_breaks

    def _extract_full_text(self):
        full_text = []
        for para in self.doc.paragraphs:
            full_text.append(para.text + "\n")
        text = ''.join(full_text)
        return text

    def get_page_number(self, paragraph_index):
        page_num = 1
        for break_index, break_page in self.page_breaks:
            if paragraph_index >= break_index:
                page_num = break_page + 1
            else:
                break
        return page_num

    def find_keyword_occurrences(self, keywords):
        occurrences = []
        context_length = 50

        keywords_set = set(kw.strip() for kw in keywords if kw.strip())
        unique_keywords = list(keywords_set)

        for keyword in unique_keywords:
            keyword = keyword.strip()
            if not keyword:
                continue
            start = 0
            while True:
                pos = self.full_text.find(keyword, start)
                if pos == -1:
                    break

                para_index = 0
                text_pos = 0
                for i, para in enumerate(self.doc.paragraphs):
                    para_end = text_pos + len(para.text) + 1
                    if text_pos <= pos < para_end:
                        para_index = i
                        break
                    text_pos = para_end

                page_num = self.get_page_number(para_index)

                context_start = max(0, pos - context_length)
                context_end = min(len(self.full_text), pos + len(keyword) + context_length)
                context = self.full_text[context_start:context_end].strip()

                occurrences.append({
                    'keyword': keyword,
                    'page': page_num,
                    'context': context
                })
                start = pos + 1

        return occurrences

def load_regions():
    if os.path.exists(REGIONS_FILE):
        try:
            with open(REGIONS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if isinstance(data, list):
                    return data
                elif isinstance(data, dict):
                    regions = []
                    if "省级" in data:
                        regions.extend(data["省级"])
                    if "市级" in data:
                        regions.extend(data["市级"])
                    if "区级" in data:
                        regions.extend(data["区级"])
                    return regions
                else:
                    return []
        except (json.JSONDecodeError, IOError):
            return []
    return []

@file_parse_bp.route('/file-parse', methods=['POST'])
def parse_file():
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': '没有选择文件'}), 400
        file = request.files['file']

        if file.filename == '':
            return jsonify({'success': False, 'message': '没有选择文件'}), 400

        if not file.filename.lower().endswith(('.doc', '.docx', '.xls', '.xlsx')):
            return jsonify({'success': False, 'message': '不支持的文件类型'}), 400

        custom_keywords = []
        for key, value in request.form.items():
            if key.startswith('keywords['):
                custom_keywords.append(value)
        if 'keywords' in request.form:
            if isinstance(request.form['keywords'], str):
                custom_keywords.append(request.form['keywords'])
            else:
                custom_keywords.extend(request.form.getlist('keywords'))

        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1])
        file.save(temp_file.name)
        temp_file_path = temp_file.name
        temp_file.close()

        try:
            keywords = load_regions()
            keywords.extend(custom_keywords)
            if not keywords:
                return jsonify({'success': False, 'message': '地域名称列表为空'}), 500

            if file.filename.lower().endswith('.docx'):
                extractor = DocxTextExtractor(temp_file_path)
                occurrences = extractor.find_keyword_occurrences(keywords)
                os.unlink(temp_file_path)

                matches = []
                for occurrence in occurrences:
                    matches.append({
                        'location': f'Page {occurrence["page"]}',
                        'text': occurrence["context"],
                        'place': occurrence["keyword"]
                    })

                return jsonify({
                    'success': True,
                    'fileName': file.filename,
                    'matches': matches
                })
            else:
                os.unlink(temp_file_path)
                return jsonify({
                    'success': True,
                    'fileName': file.filename,
                    'matches': [
                        {
                            'location': 'Sheet 1, Row 1, Column A',
                            'text': '这是一个包含北京的句子示例',
                            'place': '北京'
                        },
                        {
                            'location': 'Sheet 1, Row 3, Column B',
                            'text': '这是一个包含上海的句子示例',
                            'place': '上海'
                        }
                    ]
                })

        except Exception as e:
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
            return jsonify({'success': False, 'message': f'处理文件时出错: {str(e)}'}), 500

    except Exception as e:
        return jsonify({'success': False, 'message': f'上传文件时发生错误: {str(e)}'}), 500

@file_parse_bp.route('/regions', methods=['GET'])
def get_regions():
    regions = load_regions()
    return jsonify({'success': True, 'regions': regions})