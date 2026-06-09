"""
文档解析模块
支持 PDF、DOCX、图片等文件的文本提取
"""
import os
import re
import json
from datetime import datetime


def parse_file(file_path, file_type):
    """
    解析文件，提取文本内容
    返回: {
        "text": "完整文本",
        "pages": [{"page_no": 1, "text": "..."}],
        "tables": [],
        "page_count": 10
    }
    """
    result = {
        "text": "",
        "pages": [],
        "tables": [],
        "page_count": 0,
        "error": None
    }

    try:
        if file_type == "PDF":
            result = _parse_pdf(file_path)
        elif file_type in ["DOC", "DOCX"]:
            result = _parse_docx(file_path)
        elif file_type == "IMAGE":
            result = _parse_image(file_path)
        elif file_type == "TEXT":
            result = _parse_text(file_path)
        elif file_type == "HTML":
            result = _parse_html(file_path)
        else:
            result["error"] = f"不支持的文件类型: {file_type}"
    except Exception as e:
        result["error"] = str(e)

    return result


def _parse_pdf(file_path):
    """解析 PDF 文件"""
    result = {"text": "", "pages": [], "tables": [], "page_count": 0, "error": None}

    try:
        import pdfplumber
    except ImportError:
        result["error"] = "请先安装 pdfplumber: pip install pdfplumber"
        return result

    try:
        with pdfplumber.open(file_path) as pdf:
            result["page_count"] = len(pdf.pages)
            all_text = []

            for i, page in enumerate(pdf.pages):
                page_no = i + 1
                page_text = page.extract_text() or ""

                # 提取表格
                tables = page.extract_tables()
                for table in tables:
                    table_text = _table_to_text(table)
                    page_text += "\n" + table_text

                all_text.append(page_text)
                result["pages"].append({
                    "page_no": page_no,
                    "text": page_text,
                    "tables": tables
                })

            result["text"] = "\n\n".join(all_text)
    except Exception as e:
        result["error"] = f"PDF 解析失败: {str(e)}"

    return result


def _parse_docx(file_path):
    """解析 DOCX 文件"""
    result = {"text": "", "pages": [], "tables": [], "page_count": 1, "error": None}

    try:
        from docx import Document
    except ImportError:
        result["error"] = "请先安装 python-docx: pip install python-docx"
        return result

    try:
        doc = Document(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # 提取表格
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
                paragraphs.append(" | ".join(row_data))
            result["tables"].append(table_data)

        result["text"] = "\n".join(paragraphs)
        result["pages"] = [{"page_no": 1, "text": result["text"]}]
    except Exception as e:
        result["error"] = f"DOCX 解析失败: {str(e)}"

    return result


def _parse_image(file_path):
    """解析图片文件（OCR）"""
    result = {"text": "", "pages": [], "tables": [], "page_count": 1, "error": None}

    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        result["error"] = "请先安装 pytesseract 和 Pillow: pip install pytesseract Pillow"
        return result

    try:
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img, lang="chi_sim+eng")
        result["text"] = text.strip()
        result["pages"] = [{"page_no": 1, "text": result["text"]}]
    except Exception as e:
        result["error"] = f"图片 OCR 失败: {str(e)}"

    return result


def _parse_text(file_path):
    """解析纯文本文件"""
    result = {"text": "", "pages": [], "tables": [], "page_count": 1, "error": None}

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            result["text"] = f.read()
        result["pages"] = [{"page_no": 1, "text": result["text"]}]
    except Exception as e:
        result["error"] = f"文本文件读取失败: {str(e)}"

    return result


def _parse_html(file_path):
    """解析 HTML 文件"""
    result = {"text": "", "pages": [], "tables": [], "page_count": 1, "error": None}

    try:
        from bs4 import BeautifulSoup
    except ImportError:
        result["error"] = "请先安装 beautifulsoup4: pip install beautifulsoup4"
        return result

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, "html.parser")

        # 移除脚本和样式
        for script in soup(["script", "style"]):
            script.decompose()

        result["text"] = soup.get_text(separator="\n", strip=True)
        result["pages"] = [{"page_no": 1, "text": result["text"]}]
    except Exception as e:
        result["error"] = f"HTML 解析失败: {str(e)}"

    return result


def _table_to_text(table):
    """将表格转换为文本"""
    if not table:
        return ""

    lines = []
    for row in table:
        if row:
            line = " | ".join([str(cell) if cell else "" for cell in row])
            lines.append(line)

    return "\n".join(lines)


def extract_text_segments(text, max_length=4000):
    """
    将长文本分割成多个片段，用于 AI 处理
    """
    if len(text) <= max_length:
        return [text]

    segments = []
    paragraphs = text.split("\n\n")
    current_segment = ""

    for para in paragraphs:
        if len(current_segment) + len(para) + 2 <= max_length:
            current_segment += "\n\n" + para if current_segment else para
        else:
            if current_segment:
                segments.append(current_segment)
            current_segment = para

    if current_segment:
        segments.append(current_segment)

    return segments


def find_risk_paragraphs(text, keywords=None, special_marks=None):
    """
    查找文本中可能包含风险条款的段落
    """
    if keywords is None:
        keywords = ["废标", "否决投标", "无效投标", "投标无效", "实质性", "不得负偏离",
                    "必须满足", "资格审查", "符合性审查", "星号条款"]
    if special_marks is None:
        special_marks = ["★", "*", "▲", "#", "※"]

    paragraphs = text.split("\n")
    risk_paragraphs = []

    for i, para in enumerate(paragraphs):
        para = para.strip()
        if not para:
            continue

        # 检查关键词
        matched_keywords = [kw for kw in keywords if kw in para]

        # 检查特殊标记
        matched_marks = [m for m in special_marks if m in para]

        if matched_keywords or matched_marks:
            risk_paragraphs.append({
                "index": i,
                "text": para,
                "keywords": matched_keywords,
                "marks": matched_marks,
            })

    return risk_paragraphs


def find_special_mark_definitions(text, special_marks=None):
    """
    查找特殊符号的定义说明
    例如："带 ★ 号条款为实质性条款，不满足作无效投标处理"
    """
    if special_marks is None:
        special_marks = ["★", "*", "▲", "#", "※"]

    definitions = {}
    patterns = [
        r"带\s*([★*▲#※])\s*[号条款]*.{0,50}(实质性|无效投标|废标|否决|不得负偏离)",
        r"[标]注\s*([★*▲#※])\s*[的号]*.{0,50}(实质性|重要|不得负偏离)",
        r"([★*▲#※])\s*[号条款]*.{0,50}(实质性|无效投标|废标|否决)",
    ]

    for mark in special_marks:
        for pattern in patterns:
            pattern_with_mark = pattern.replace("([★*▲#※])", f"({re.escape(mark)})")
            matches = re.finditer(pattern_with_mark, text)
            for match in matches:
                if mark not in definitions:
                    definitions[mark] = []
                definitions[mark].append({
                    "text": match.group(0),
                    "is_risk": any(kw in match.group(0) for kw in ["废标", "否决", "无效投标", "实质性", "不得负偏离"])
                })

    return definitions
