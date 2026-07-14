"""
投标比对 — 文件检查 API Blueprint
- 选择已分析的招标项目（提供招标文件内容），上传投标响应文件（docx）
- AI 对照检查缺漏与风险，生成 docx 结果报告供下载
- 结果不持久化：仅本次请求返回，刷新后清空
"""
import os
import io
import json
import tempfile
from datetime import datetime
from flask import Blueprint, jsonify, request, send_file, Response

bid_check_bp = Blueprint("bid_check", __name__)

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DATA_DIR = os.path.join(BASE_DIR, "data")
PARSED_DIR = os.path.join(DATA_DIR, "parsed")
PROJECTS_FILE = os.path.join(DATA_DIR, "projects.json")

MAX_BID_FILE_SIZE = 200 * 1024 * 1024


def _load_projects():
    if os.path.exists(PROJECTS_FILE):
        try:
            with open(PROJECTS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return []
    return []


def _get_tender_content(project_id):
    parts = []
    project_parsed_dir = os.path.join(PARSED_DIR, project_id)

    md_path = os.path.join(project_parsed_dir, "announcement.md")
    if os.path.exists(md_path):
        try:
            with open(md_path, "r", encoding="utf-8") as f:
                md = f.read().strip()
            if md:
                parts.append("【公告网页内容】\n" + md)
        except Exception:
            pass

    if os.path.isdir(project_parsed_dir):
        for fname in sorted(os.listdir(project_parsed_dir)):
            if not fname.endswith(".json"):
                continue
            fpath = os.path.join(project_parsed_dir, fname)
            try:
                with open(fpath, "r", encoding="utf-8") as f:
                    data = json.load(f)
            except Exception:
                continue
            text = (data.get("text") or "").strip()
            tables = data.get("tables") or []
            table_parts = []
            for idx, table in enumerate(tables):
                rows = []
                for row in table:
                    if isinstance(row, list):
                        rows.append(" | ".join(str(c) for c in row))
                    else:
                        rows.append(str(row))
                if rows:
                    table_parts.append("[表格" + str(idx + 1) + "]\n" + "\n".join(rows))
            if table_parts:
                text = (text + "\n\n" + "\n\n".join(table_parts)).strip()
            if text:
                parts.append("【招标文件解析内容】\n" + text)

    return "\n\n".join(parts)


def _truncate(text, max_chars):
    if not text or len(text) <= max_chars:
        return text or ""
    head = int(max_chars * 0.75)
    tail = max_chars - head
    return (
        text[:head]
        + "\n\n…（内容过长，已截取开头与结尾部分，完整内容请参见原文件）…\n\n"
        + text[-tail:]
    )


def _generate_bid_check_stream(messages, model_id, project, original_name, bid_truncated, tender_truncated):
    """生成流式响应的生成器"""
    from backend.ai_extractor import _parse_json_response
    from backend.llm_client import get_all_enabled_model_configs, call_llm

    try:
        target_model = None
        if model_id:
            all_models = get_all_enabled_model_configs()
            target_model = next((m for m in all_models if str(m.get("id")) == str(model_id)), None)

        if target_model:
            model_name = target_model.get("name", target_model.get("model", "unknown"))
            stream_gen = call_llm(messages, temperature=0.1, max_tokens=8000, model_config=target_model, stream=True)
        else:
            stream_gen = call_llm(messages, temperature=0.1, max_tokens=8000, stream=True)
            model_name = ""

        yield json.dumps({"type": "status", "message": "正在分析..."}, ensure_ascii=False) + "\n"

        full_response = ""
        for chunk in stream_gen:
            full_response += chunk
            yield json.dumps({"type": "stream", "chunk": chunk}, ensure_ascii=False) + "\n"

        yield json.dumps({"type": "status", "message": "正在解析结果..."}, ensure_ascii=False) + "\n"

        parsed = _parse_json_response(full_response)
        if not parsed or not isinstance(parsed, dict):
            yield json.dumps({
                "type": "error",
                "message": "AI 返回格式解析失败，请稍后重试或调整提示词",
                "available_models": [{"id": m.get("id"), "name": m.get("name")} for m in get_all_enabled_model_configs()]
            }, ensure_ascii=False) + "\n"
            return

        summary = str(parsed.get("summary", "")).strip()
        raw_items = parsed.get("items") or []
        valid_cats = {"基本信息", "资格项", "完整性", "废标项", "商务和技术文件", "得分项"}
        valid_risk = {"高", "中", "低", "符合"}
        norm_items = []
        for it in raw_items:
            if not isinstance(it, dict):
                continue
            risk = str(it.get("risk_level", "")).strip()
            if risk not in valid_risk:
                risk = "低"
            cat = str(it.get("category", "")).strip()
            if cat not in valid_cats:
                cat = "基本信息"
            norm_items.append({
                "category": cat,
                "risk_level": risk,
                "title": str(it.get("title", "")).strip() or "—",
                "detail": str(it.get("detail", "")).strip(),
                "suggestion": str(it.get("suggestion", "")).strip(),
            })

        yield json.dumps({
            "type": "complete",
            "data": {
                "success": True,
                "project_name": project.get("name", ""),
                "bid_file_name": original_name,
                "checked_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "summary": summary or "检查完成。",
                "items": norm_items,
                "truncated": bool(bid_truncated or tender_truncated),
                "model": model_name,
            }
        }, ensure_ascii=False) + "\n"

    except Exception as e:
        yield json.dumps({
            "type": "error",
            "message": "AI 调用失败：" + str(e),
            "available_models": [{"id": m.get("id"), "name": m.get("name")} for m in get_all_enabled_model_configs()]
        }, ensure_ascii=False) + "\n"


# ── 运行检查 ───────────────────────────────────────────────
@bid_check_bp.route("/api/bid-check/run", methods=["POST"])
def run_bid_check():
    project_id = (request.form.get("project_id") or "").strip()
    prompt = (request.form.get("prompt") or "").strip()
    model_id = request.form.get('model_id')

    if "file" not in request.files:
        return jsonify({"success": False, "message": "请上传投标响应文件（docx）"}), 400
    f = request.files["file"]
    if not f.filename:
        return jsonify({"success": False, "message": "文件名为空"}), 400
    if not project_id:
        return jsonify({"success": False, "message": "请选择招标项目"}), 400
    if not prompt:
        return jsonify({"success": False, "message": "检查提示词不能为空"}), 400

    original_name = f.filename
    ext = os.path.splitext(original_name)[1].lower()
    if ext != ".docx":
        return jsonify({"success": False, "message": "投标响应文件必须为 .docx 格式"}), 400

    cl = request.content_length or 0
    if cl > MAX_BID_FILE_SIZE + 10 * 1024 * 1024:
        return jsonify({"success": False, "message": "文件大小超过 200MB 限制"}), 413

    projects = _load_projects()
    project = next((p for p in projects if p.get("id") == project_id), None)
    if not project:
        return jsonify({"success": False, "message": "所选招标项目不存在"}), 404

    tender_content = _get_tender_content(project_id)
    if not tender_content.strip():
        return jsonify({
            "success": False,
            "message": "该招标项目尚未解析任何招标文件内容，请先在「招标分析」中对该项目执行分析"
        }), 400

    import sys
    sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    from backend.doc_parser import parse_file

    bid_text = ""
    bid_truncated = False
    with tempfile.TemporaryDirectory() as td:
        tmp_path = os.path.join(td, "bid_upload" + ext)
        f.save(tmp_path)
        if os.path.getsize(tmp_path) > MAX_BID_FILE_SIZE:
            return jsonify({"success": False, "message": "文件大小超过 200MB 限制"}), 413

        result = parse_file(tmp_path, "DOCX")
        if result.get("error"):
            return jsonify({"success": False, "message": "投标文件解析失败：" + result["error"]}), 400

        text = (result.get("text") or "").strip()
        tables = result.get("tables") or []
        table_parts = []
        for idx, table in enumerate(tables):
            rows = []
            for row in table:
                if isinstance(row, list):
                    rows.append(" | ".join(str(c) for c in row))
                else:
                    rows.append(str(row))
            if rows:
                table_parts.append("[表格" + str(idx + 1) + "]\n" + "\n".join(rows))
        if table_parts:
            text = (text + "\n\n" + "\n\n".join(table_parts)).strip()

        if len(text) > 140000:
            text = _truncate(text, 140000)
            bid_truncated = True
        bid_text = text

    tender_truncated = False
    if len(tender_content) > 80000:
        tender_content = _truncate(tender_content, 80000)
        tender_truncated = True

    system_prompt = (
        "你是一个专业的投标文件检查助手。请严格按照用户给出的检查要求，对照【招标文件内容】与"
        "【投标响应文件内容】逐项核查，并以严格的 JSON 格式输出结果。"
        "只返回 JSON 对象本身，不要包含 ``` 代码块或任何解释文字。\n\n"
        "输出格式：\n"
        "{\n"
        '  "summary": "对投标响应文件整体符合性的总体评估（150字以内）",\n'
        '  "items": [\n'
        '    {\n'
        '      "category": "检查类别，从以下选择之一：基本信息|资格项|完整性|废标项|商务和技术文件|得分项",\n'
        '      "risk_level": "风险等级，从以下选择之一：高|中|低|符合",\n'
        '      "title": "事项或问题标题（30字以内）",\n'
        '      "detail": "详细说明，引用招标文件要求与投标文件实际内容进行对照",\n'
        '      "suggestion": "修改建议（若符合要求可留空字符串）"\n'
        '    }\n'
        "  ]\n"
        "}\n\n"
        "注意：\n"
        "1. risk_level 含义：高=可能导致废标；中=响应不完整/不一致；低=笔误/表述优化；符合=该检查项未发现问题\n"
        "2. 用户的检查要求中每一事项至少产出 1 个 item；若该事项全部符合，则产出 1 个 risk_level=符合 的 item\n"
        "3. items 必须覆盖用户提出的全部检查事项，不得遗漏\n"
        "4. detail 中需指明该结论依据招标文件与投标文件中的哪些内容\n"
        "5. 始终以采购文件为准进行检查，不自行推断要求"
    )

    user_content = (
        prompt.strip()
        + "\n\n==== 招标文件内容 ====\n" + tender_content
        + "\n\n==== 投标响应文件内容 ====\n"
        + (bid_text or "（投标响应文件未提取到文本内容，可能为扫描件或空文件）")
    )

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_content},
    ]

    return Response(
        _generate_bid_check_stream(messages, model_id, project, original_name, bid_truncated, tender_truncated),
        content_type="text/event-stream"
    )


# ── 下载结果报告（docx） ──────────────────────────────────
@bid_check_bp.route("/api/bid-check/download", methods=["POST"])
def download_bid_check_report():
    data = request.get_json(silent=True) or {}
    if not data.get("items") and not data.get("summary"):
        return jsonify({"success": False, "message": "报告内容为空"}), 400

    try:
        bio = _build_report_docx(data)
    except Exception as e:
        return jsonify({"success": False, "message": "报告生成失败：" + str(e)}), 500

    fname = f"投标文件检查报告_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    return send_file(bio, as_attachment=True, download_name=fname,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def _build_report_docx(data):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    normal = doc.styles['Normal']
    normal.font.name = '宋体'
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    title = doc.add_heading('投标文件检查报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run('招标项目：').bold = True
    meta.add_run(data.get('project_name', '—') + '\n')
    meta.add_run('投标文件：').bold = True
    meta.add_run(data.get('bid_file_name', '—') + '\n')
    meta.add_run('检查时间：').bold = True
    meta.add_run(data.get('checked_at', '—'))
    for run in meta.runs:
        run.font.size = Pt(10)

    doc.add_paragraph('')

    doc.add_heading('一、总体评估', level=1)
    doc.add_paragraph(data.get('summary', '—'))

    if data.get('truncated'):
        note_p = doc.add_paragraph()
        rn = note_p.add_run('注：因文件篇幅较大，已对招标/投标文件内容做截取处理，部分内容未被纳入 AI 检查。')
        rn.italic = True
        rn.font.size = Pt(9)
        rn.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_heading('二、检查明细', level=1)
    risk_label = {"高": "高风险", "中": "中风险", "低": "低风险", "符合": "符合"}
    risk_color = {
        "高": RGBColor(0xC0, 0x39, 0x2B),
        "中": RGBColor(0xCC, 0x7A, 0x00),
        "低": RGBColor(0xB8, 0x86, 0x00),
        "符合": RGBColor(0x2E, 0x7D, 0x32),
    }
    items = data.get('items') or []

    if not items:
        doc.add_paragraph('无检查项。')
    else:
        for idx, it in enumerate(items, 1):
            risk = it.get('risk_level', '低')
            p = doc.add_paragraph()
            run = p.add_run(f"{idx}. [{risk_label.get(risk, risk)}] {it.get('title', '—')}")
            run.bold = True
            run.font.color.rgb = risk_color.get(risk, RGBColor(0x33, 0x33, 0x33))

            p2 = doc.add_paragraph()
            r2 = p2.add_run('类别：' + it.get('category', '—'))
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            doc.add_paragraph('说明：' + it.get('detail', '—'))
            if it.get('suggestion'):
                sug = doc.add_paragraph()
                rs = sug.add_run('建议：' + it['suggestion'])
                rs.italic = True
                rs.font.size = Pt(10)
            doc.add_paragraph('')

    footer = doc.add_paragraph()
    rn = footer.add_run('本报告由 BidTool 投标工具 自动生成，仅供评审参考。')
    rn.font.size = Pt(9)
    rn.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    rn.italic = True

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
