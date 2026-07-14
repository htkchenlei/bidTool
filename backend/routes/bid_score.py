"""
模拟评分 API Blueprint（新版）
- 上传招标文件，使用 markitdown 转换为 markdown
- AI 提取评分标准，以表格形式展示
- 上传一份投标文件进行模拟评分
- 双轨评分：图片证据验证 + 文本内容匹配
- 结果不持久化：仅本次请求返回，刷新后清空
"""
import os
import io
import re
import json
import tempfile
import traceback
from datetime import datetime
from flask import Blueprint, jsonify, request, Response, send_file

bid_score_bp = Blueprint("bid_score", __name__)

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DATA_DIR = os.path.join(BASE_DIR, "data")
MAX_BID_FILE_SIZE = 200 * 1024 * 1024

EXT_TYPE_MAP = {
    ".docx": "DOCX",
    ".doc": "DOCX",
    ".pdf": "PDF",
    ".png": "IMAGE",
    ".jpg": "IMAGE",
    ".jpeg": "IMAGE",
    ".txt": "TEXT",
}


def _truncate(text, max_chars):
    if not text or len(text) <= max_chars:
        return text or ""
    head = int(max_chars * 0.75)
    tail = max_chars - head
    return (
        text[:head]
        + "\n\n…（内容过长，已截取开头与结尾部分，完整内容请参见原文件）…\n\n"
        + text[-tail:]
    )


def _parse_json_response(response):
    default = {}
    text = response.strip()
    if "```json" in text:
        start = text.index("```json") + 7
        end = text.index("```", start)
        text = text[start:end].strip()
    elif "```" in text:
        start = text.index("```") + 3
        end = text.index("```", start)
        text = text[start:end].strip()
    try:
        data = json.loads(text)
        return data
    except json.JSONDecodeError:
        try:
            data = json.loads(text.replace("'", '"'))
            return data
        except:
            pass
        result = {}
        patterns = {
            "criteria": r'"criteria"\s*:\s*\[.*?\]',
            "total_max": r'"total_max"\s*:\s*[\d.]+',
            "price_rule": r'"price_rule"\s*:\s*\{.*?\}',
            "score": r'"score"\s*:\s*[\d.]+',
            "max_score": r'"max_score"\s*:\s*[\d.]+',
            "reason": r'"reason"\s*:\s*"[^"]*"',
            "evidence": r'"evidence"\s*:\s*"[^"]*"',
            "confidence": r'"confidence"\s*:\s*[\d.]+',
            "cert_name": r'"cert_name"\s*:\s*"[^"]*"',
            "cert_type": r'"cert_type"\s*:\s*"[^"]*"',
            "issuer": r'"issuer"\s*:\s*"[^"]*"',
            "expire_date": r'"expire_date"\s*:\s*"[^"]*"',
            "is_valid": r'"is_valid"\s*:\s*(true|false)',
        }
        for key, pat in patterns.items():
            m = re.search(pat, response)
            if m:
                try:
                    if key in ["score", "max_score", "confidence"]:
                        result[key] = float(m.group(0).split(":")[1].strip())
                    elif key == "is_valid":
                        result[key] = m.group(0).split(":")[1].strip().lower() == "true"
                    elif key in ["criteria", "price_rule"]:
                        result[key] = json.loads("{" + m.group(0) + "}")[key]
                    else:
                        result[key] = m.group(0).split(":")[1].strip().strip('"')
                except:
                    pass
        return result


def _extract_criteria_section(tender_content):
    lines = tender_content.split('\n')
    start_idx = -1
    end_idx = -1
    
    for i, line in enumerate(lines):
        if '项目 | 评审因素' in line and '分值' in line:
            start_idx = max(0, i - 2)
            break
    
    if start_idx == -1:
        for i, line in enumerate(lines):
            if '评分标准' in line and len(line) < 50:
                start_idx = max(0, i - 2)
                break
    
    if start_idx == -1:
        for i, line in enumerate(lines):
            if '第五章' in line and ('评标' in line or '评分' in line):
                start_idx = max(0, i - 2)
                break
    
    if start_idx >= 0:
        table_count = 0
        for i in range(start_idx, min(start_idx + 600, len(lines))):
            line = lines[i]
            if '|' in line and ('项目' in line or '评审因素' in line or '分值' in line):
                table_count += 1
            
            if table_count >= 3 and line.strip() and '|' not in line and i > start_idx + 50:
                if '第七章' in line or '第六章' in line or '投标报价合计' in line:
                    end_idx = i
                    break
            
            if i == start_idx + 599:
                end_idx = i + 1
                break
    
    if start_idx >= 0 and end_idx > start_idx:
        return '\n'.join(lines[start_idx:end_idx])
    
    for i, line in enumerate(lines):
        if '项目 | 评审因素' in line:
            return '\n'.join(lines[max(0, i-2):min(len(lines), i+150)])
    
    return tender_content


def _build_criteria_extraction_messages(user_prompt, tender_content):
    criteria_section = _extract_criteria_section(tender_content)
    
    system_prompt = (
        "你是一位专业的政府采购评审专家。请从招标文件的评分标准表格中提取各项评分要求。\n\n"
        "严格按照以下格式输出，只返回JSON，不要包含任何其他文字、解释或思考过程：\n\n"
        "{\n"
        '  "criteria": [\n'
        '    {\n'
        '      "id": "c1",\n'
        '      "name": "投标报价",\n'
        '      "max_score": 10,\n'
        '      "category": "price",\n'
        '      "requires_image": false,\n'
        '      "description": "报价分采用低价优先法计算",\n'
        '      "sub_items": [],\n'
        '      "keywords": ["报价", "价格"]\n'
        '    }\n'
        '  ],\n'
        '  "total_max": 100,\n'
        '  "price_rule": {"method": "lowest", "price_weight": 10}\n'
        "}\n\n"
        "规则：\n"
        "1. category: price(价格部分), commercial(商务部分如业绩、资质证书、团队人员), technical(技术部分如方案、制度、措施)\n"
        "2. requires_image: 需要图片证据的为true（如合同扫描件、资质证书、人员证书），不需要的为false\n"
        "3. keywords: 提取2-5个关键词"
    )

    user_content = (
        user_prompt
        + "\n\n==== 评分标准内容 ====\n" + _truncate(criteria_section, 40000)
    )

    return [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_content},
    ]


def _normalize_criteria(parsed):
    if not parsed or not isinstance(parsed, dict):
        return None

    criteria = parsed.get("criteria") or []
    if not isinstance(criteria, list) or not criteria:
        return None

    norm_criteria = []
    for idx, c in enumerate(criteria):
        if not isinstance(c, dict):
            continue
        name = str(c.get("name", "")).strip()
        if not name:
            continue
        try:
            max_score = float(c.get("max_score", 0) or 0)
        except Exception:
            max_score = 0
        category = str(c.get("category", "")).strip().lower()
        if category not in ("price", "commercial", "technical"):
            if any(kw in name for kw in ["价", "报价", "价格"]):
                category = "price"
            elif any(kw in name for kw in ["资质", "业绩", "证书", "社保", "案例", "人员", "团队"]):
                category = "commercial"
            else:
                category = "technical"
        requires_image = bool(c.get("requires_image", False))
        description = str(c.get("description", "")).strip()
        keywords = []
        raw_kw = c.get("keywords") or []
        if isinstance(raw_kw, list):
            keywords = [str(x).strip() for x in raw_kw if str(x).strip()]
        sub_items = []
        raw_sub = c.get("sub_items") or []
        if isinstance(raw_sub, list):
            for si in raw_sub:
                if not isinstance(si, dict):
                    continue
                si_name = str(si.get("name", "")).strip()
                if not si_name:
                    continue
                try:
                    si_max = float(si.get("max_score", 0) or 0)
                except Exception:
                    si_max = 0
                si_ri = bool(si.get("requires_image", False))
                sub_items.append({
                    "name": si_name,
                    "max_score": si_max,
                    "requires_image": si_ri,
                })

        norm_criteria.append({
            "id": str(idx + 1),
            "name": name,
            "max_score": max_score,
            "category": category,
            "requires_image": requires_image,
            "description": description,
            "keywords": keywords,
            "sub_items": sub_items,
        })

    if not norm_criteria:
        return None

    total_max = 0
    try:
        total_max = float(parsed.get("total_max", sum(c["max_score"] for c in norm_criteria)) or 0)
    except Exception:
        total_max = sum(c["max_score"] for c in norm_criteria)

    price_rule = None
    has_price = any(c["category"] == "price" for c in norm_criteria)
    if has_price:
        price_weight = sum(c["max_score"] for c in norm_criteria if c["category"] == "price")
        price_rule = {
            "method": "base_price",
            "price_weight": price_weight,
        }

    return {
        "criteria": norm_criteria,
        "total_max": total_max,
        "price_rule": price_rule,
    }


def _generate_strengths_weaknesses(scores):
    strengths = []
    weaknesses = []

    for score in scores:
        pct = score["score"] / score["max_score"] if score["max_score"] > 0 else 0
        if pct >= 0.8:
            strengths.append(f"{score['name']}得分较高（{score['score']}/{score['max_score']}）")
        elif pct < 0.5:
            weaknesses.append(f"{score['name']}得分较低（{score['score']}/{score['max_score']}）")

    return strengths, weaknesses


def _adjust_scores_by_comparison(bid_results, criteria):
    bid_count = len(bid_results)
    if bid_count < 2:
        return bid_results

    grade_ranges = {
        2: [(0.90, 1.00), (0.50, 0.89)],
        3: [(0.90, 1.00), (0.65, 0.89), (0.30, 0.64)],
        4: [(0.90, 1.00), (0.75, 0.89), (0.50, 0.74), (0.20, 0.49)],
        5: [(0.90, 1.00), (0.75, 0.89), (0.60, 0.74), (0.40, 0.59), (0.00, 0.39)],
    }
    ranges = grade_ranges.get(min(bid_count, 5), grade_ranges[5])

    for criterion in criteria:
        criterion_name = criterion["name"]
        max_score = criterion["max_score"]

        scores_with_bid = []
        for bid_idx, bid_result in enumerate(bid_results):
            for score in bid_result["scores"]:
                if score["name"] == criterion_name:
                    scores_with_bid.append({
                        "bid_idx": bid_idx,
                        "original_score": score["score"],
                        "reason": score.get("reason", ""),
                        "evidence": score.get("evidence", ""),
                    })
                    break

        scores_with_bid.sort(key=lambda x: x["original_score"], reverse=True)

        for rank, item in enumerate(scores_with_bid):
            grade_idx = min(rank, len(ranges) - 1)
            lower, upper = ranges[grade_idx]

            adjusted_score = round(max_score * (lower + upper) / 2, 2)

            bid_result = bid_results[item["bid_idx"]]
            for score in bid_result["scores"]:
                if score["name"] == criterion_name:
                    old_score = score["score"]
                    score["score"] = adjusted_score
                    score["reason"] = f"[对比调整] {item['reason'] or '综合评定'}（原得分{old_score}，排名第{rank+1}，调整为{adjusted_score}）"
                    break

    for bid_result in bid_results:
        bid_result["total_score"] = round(sum(s["score"] for s in bid_result["scores"]), 2)

    return bid_results


DEFAULT_CRITERIA = [
    {"id": "1", "name": "技术方案", "max_score": 30, "category": "technical", "requires_image": False, "description": "技术方案完整性与可行性", "keywords": ["技术方案", "实施方案"], "sub_items": []},
    {"id": "2", "name": "商务报价", "max_score": 20, "category": "price", "requires_image": False, "description": "投标报价合理性", "keywords": ["报价", "价格"], "sub_items": []},
    {"id": "3", "name": "企业资质", "max_score": 15, "category": "commercial", "requires_image": True, "description": "企业资质证书", "keywords": ["资质", "证书", "营业执照"], "sub_items": []},
    {"id": "4", "name": "项目业绩", "max_score": 15, "category": "commercial", "requires_image": False, "description": "类似项目业绩", "keywords": ["业绩", "案例", "项目经验"], "sub_items": []},
    {"id": "5", "name": "团队资质", "max_score": 10, "category": "commercial", "requires_image": True, "description": "项目团队人员资质", "keywords": ["团队", "人员", "职称"], "sub_items": []},
    {"id": "6", "name": "服务承诺", "max_score": 10, "category": "technical", "requires_image": False, "description": "服务承诺与保障", "keywords": ["服务", "承诺"], "sub_items": []},
]
DEFAULT_TOTAL_MAX = 100


@bid_score_bp.route("/api/bid-score/extract-criteria", methods=["POST"])
def extract_criteria():
    from backend.markdown_converter import convert_to_markdown, detect_file_type
    from backend.llm_client import call_llm, get_all_enabled_model_configs

    project_id = request.form.get("project_id")
    file = request.files.get("file")
    model_id = request.form.get("model_id")
    user_prompt = (request.form.get("prompt") or "").strip()

    tender_content = ""
    tender_file_name = ""

    if project_id:
        from backend.routes.projects import load_project_files
        project_files = load_project_files()
        tender_files = [f for f in project_files 
                        if f.get("project_id") == project_id 
                        and f.get("storage_name", "").lower().endswith((".docx", ".pdf"))]
        
        if not tender_files:
            return jsonify({"success": False, "message": "该项目没有招标文件，请先上传"}), 400

        tender_file = tender_files[0]
        tender_file_name = tender_file.get("original_name", "")
        storage_path = tender_file.get("storage_path", "")
        
        if not os.path.exists(storage_path):
            return jsonify({"success": False, "message": "招标文件不存在"}), 400

        file_type = detect_file_type(storage_path)
        md_result = convert_to_markdown(storage_path, file_type)
        
        if not md_result["success"]:
            return jsonify({"success": False, "message": f"文档转换失败：{md_result.get('error', '未知错误')}"}), 400
        
        tender_content = md_result["markdown"]

    elif file and file.filename:
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in (".docx", ".pdf"):
            return jsonify({"success": False, "message": "招标文件仅支持 .docx / .pdf 格式"}), 400

        with tempfile.TemporaryDirectory() as td:
            safe_name = f"tender{ext}"
            tmp_path = os.path.join(td, safe_name)
            file.save(tmp_path)

            file_type = detect_file_type(tmp_path)
            md_result = convert_to_markdown(tmp_path, file_type)

            if not md_result["success"]:
                return jsonify({"success": False, "message": f"文档转换失败：{md_result.get('error', '未知错误')}"}), 400

            tender_content = md_result["markdown"]
            tender_file_name = file.filename
    else:
        return jsonify({"success": False, "message": "请上传招标文件或选择项目"}), 400

    if not tender_content.strip():
        return jsonify({"success": False, "message": "文档转换结果为空"}), 400

    target_model = None
    if model_id:
        all_models = get_all_enabled_model_configs()
        target_model = next((m for m in all_models if str(m.get("id")) == str(model_id)), None)

    prompt_text = user_prompt if user_prompt else "请提取招标文件中的评分标准"
    messages = _build_criteria_extraction_messages(prompt_text, tender_content)

    try:
        if target_model:
            response = call_llm(messages, temperature=0.1, max_tokens=4000, model_config=target_model)
            model_name = target_model.get("name", target_model.get("model", "unknown"))
        else:
            from backend.ai_extractor import _call_llm_with_fallback
            response, model_name = _call_llm_with_fallback(messages, temperature=0.1, max_tokens=4000)
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"AI 提取失败：{str(e)}",
            "available_models": [{"id": m.get("id"), "name": m.get("name")} for m in get_all_enabled_model_configs()]
        }), 500

    parsed = _parse_json_response(response)
    criteria_info = _normalize_criteria(parsed)

    if criteria_info is None:
        criteria_info = {
            "criteria": DEFAULT_CRITERIA,
            "total_max": DEFAULT_TOTAL_MAX,
            "price_rule": None,
        }

    return jsonify({
        "success": True,
        "criteria": criteria_info["criteria"],
        "total_max": criteria_info["total_max"],
        "price_rule": criteria_info.get("price_rule"),
        "markdown_preview": tender_content[:2000],
        "model_name": model_name,
    })


@bid_score_bp.route("/api/bid-score/run", methods=["POST"])
def run_bid_score():
    project_id = request.form.get("project_id")
    tender_file = request.files.get("tender_file")
    bid_files = request.files.getlist("bid_files")
    model_id = request.form.get("model_id")
    criteria_json = request.form.get("criteria")

    if not bid_files or len(bid_files) == 0 or not bid_files[0].filename:
        return jsonify({"success": False, "message": "请上传投标文件"}), 400

    allowed_exts = (".docx", ".pdf", ".png", ".jpg", ".jpeg", ".txt")
    for bf in bid_files:
        if bf.filename:
            ext = os.path.splitext(bf.filename)[1].lower()
            if ext not in allowed_exts:
                return jsonify({"success": False, "message": f"投标文件 {bf.filename} 格式不符，仅支持 .docx / .pdf / .png / .jpg / .jpeg / .txt 格式"}), 400

    tender_path = None
    tender_file_name = ""

    if project_id:
        from backend.routes.projects import load_project_files
        project_files = load_project_files()
        tender_files = [f for f in project_files 
                        if f.get("project_id") == project_id 
                        and f.get("storage_name", "").lower().endswith((".docx", ".pdf"))]
        
        if not tender_files:
            return jsonify({"success": False, "message": "该项目没有招标文件，请先上传"}), 400

        tender_file_record = tender_files[0]
        tender_path = tender_file_record.get("storage_path", "")
        tender_file_name = tender_file_record.get("original_name", "")
        
        if not os.path.exists(tender_path):
            return jsonify({"success": False, "message": "招标文件不存在"}), 400
    elif tender_file and tender_file.filename:
        tender_ext = os.path.splitext(tender_file.filename)[1].lower()
        if tender_ext not in (".docx", ".pdf"):
            return jsonify({"success": False, "message": "招标文件仅支持 .docx / .pdf 格式"}), 400
        tender_file_name = tender_file.filename
    elif not criteria_json:
        return jsonify({"success": False, "message": "请选择项目或上传招标文件"}), 400

    criteria_info = None
    if criteria_json:
        try:
            criteria_info = json.loads(criteria_json)
        except:
            criteria_info = None

    import shutil
    upload_temp_dir = tempfile.mkdtemp(prefix="bid_score_upload_")
    
    saved_bid_files = []
    for i, bf in enumerate(bid_files):
        if bf.filename:
            ext = os.path.splitext(bf.filename)[1].lower()
            save_path = os.path.join(upload_temp_dir, f"bid_{i}{ext}")
            bf.save(save_path)
            saved_bid_files.append({
                "filename": bf.filename,
                "path": save_path
            })

    if tender_file and tender_file.filename:
        tender_ext = os.path.splitext(tender_file.filename)[1].lower()
        tender_path = os.path.join(upload_temp_dir, f"tender{tender_ext}")
        tender_file.save(tender_path)

    return Response(
        _generate_bid_score_stream(tender_file, tender_path, tender_file_name, saved_bid_files, model_id, upload_temp_dir, criteria_info),
        content_type="text/event-stream"
    )


def _generate_bid_score_stream(tender_file, tender_path, tender_file_name, bid_files, model_id, upload_temp_dir=None, criteria_info=None):
    from backend.markdown_converter import convert_to_markdown, detect_file_type
    from backend.llm_client import call_llm, get_all_enabled_model_configs, _model_supports_vision

    started_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    total_llm_call_count = 0
    current_stage = "init"

    try:
        target_model = None
        if model_id:
            all_models = get_all_enabled_model_configs()
            target_model = next((m for m in all_models if str(m.get("id")) == str(model_id)), None)

        model_name = ""
        if target_model:
            model_name = target_model.get("name", target_model.get("model", "unknown"))

        supports_vision = False
        if target_model:
            supports_vision = _model_supports_vision(target_model)

        if criteria_info is None:
            yield json.dumps({"type": "status", "stage": "parsing", "message": "正在解析招标文件...", "progress": 0.05}, ensure_ascii=False) + "\n"

            tender_tmp_path = tender_path
            if not tender_tmp_path and tender_file and tender_file.filename:
                tender_ext = os.path.splitext(tender_file.filename)[1].lower()
                tender_tmp_path = os.path.join(upload_temp_dir, f"tender{tender_ext}")
                tender_file.save(tender_tmp_path)

            tender_type = detect_file_type(tender_tmp_path)
            md_result = convert_to_markdown(tender_tmp_path, tender_type)

            if not md_result["success"]:
                yield json.dumps({
                    "type": "error",
                    "message": f"招标文件转换失败：{md_result.get('error', '未知错误')}"
                }, ensure_ascii=False) + "\n"
                return

            tender_content = md_result["markdown"]

            yield json.dumps({"type": "status", "stage": "criteria", "message": "正在提取评分标准...", "progress": 0.10}, ensure_ascii=False) + "\n"

            messages = _build_criteria_extraction_messages("请提取招标文件中的评分标准", tender_content)

            try:
                if target_model:
                    response = call_llm(messages, temperature=0.1, max_tokens=4000, model_config=target_model)
                else:
                    from backend.ai_extractor import _call_llm_with_fallback
                    response, model_name = _call_llm_with_fallback(messages, temperature=0.1, max_tokens=4000)
            except Exception as e:
                yield json.dumps({
                    "type": "error",
                    "message": f"评分标准提取失败：{str(e)}",
                    "available_models": [{"id": m.get("id"), "name": m.get("name")} for m in get_all_enabled_model_configs()]
                }, ensure_ascii=False) + "\n"
                return

            total_llm_call_count += 1
            parsed = _parse_json_response(response)
            criteria_info = _normalize_criteria(parsed)

            if criteria_info is None:
                criteria_info = {
                    "criteria": DEFAULT_CRITERIA,
                    "total_max": DEFAULT_TOTAL_MAX,
                    "price_rule": None,
                }

        criteria = criteria_info["criteria"]
        total_max = sum(c["max_score"] for c in criteria)
        bid_count = len([bf for bf in bid_files if bf.get("filename")])

        progress_start = 0.10 if criteria_info else 0.30

        all_bid_results = []

        for bid_idx, bf in enumerate(bid_files):
            if not bf.get("filename"):
                continue

            bid_filename = bf.get("filename", "")
            progress_base = progress_start + ((1.0 - progress_start) * bid_idx / max(bid_count, 1))

            yield json.dumps({
                "type": "status",
                "stage": "parsing",
                "message": f"正在解析投标文件 [{bid_idx+1}/{bid_count}]：{bid_filename}",
                "progress": round(progress_base + 0.02, 2),
                "bidder_index": bid_idx,
                "bidder_name": bid_filename,
            }, ensure_ascii=False) + "\n"

            bid_result = _score_single_bid(bf, criteria, target_model, model_name, supports_vision)

            if not bid_result["success"]:
                yield json.dumps({
                    "type": "error",
                    "message": f"投标文件 {bid_filename} 解析失败：{bid_result.get('error', '未知错误')}"
                }, ensure_ascii=False) + "\n"
                return

            total_llm_call_count += bid_result.get("llm_calls", 0)

            yield json.dumps({
                "type": "status",
                "stage": "scoring",
                "message": f"评分完成 [{bid_idx+1}/{bid_count}]：{bid_filename}，得分 {bid_result['total_score']}/{bid_result['total_max']}",
                "progress": round(progress_base + 0.08, 2),
                "bidder_index": bid_idx,
                "bidder_name": bid_filename,
            }, ensure_ascii=False) + "\n"

            all_bid_results.append(bid_result)

        yield json.dumps({"type": "status", "stage": "summary", "message": "汇总评分结果...", "progress": 0.90}, ensure_ascii=False) + "\n"

        if bid_count > 1:
            yield json.dumps({"type": "status", "stage": "comparison", "message": "正在进行投标人横向对比...", "progress": 0.93}, ensure_ascii=False) + "\n"
            all_bid_results = _adjust_scores_by_comparison(all_bid_results, criteria)

        finished_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        duration = str(datetime.strptime(finished_at, "%Y-%m-%d %H:%M:%S") - datetime.strptime(started_at, "%Y-%m-%d %H:%M:%S"))

        audit = {
            "model": model_name,
            "llm_calls": total_llm_call_count,
            "bid_count": bid_count,
            "started_at": started_at,
            "finished_at": finished_at,
            "duration": duration,
            "criteria_count": len(criteria),
            "supports_vision": supports_vision,
        }

        yield json.dumps({"type": "status", "stage": "complete", "message": "评分完成", "progress": 1.0}, ensure_ascii=False) + "\n"

        yield json.dumps({
            "type": "complete",
            "data": {
                "success": True,
                "criteria": criteria,
                "total_max": total_max,
                "bid_results": all_bid_results,
                "audit": audit,
                "timestamp": finished_at,
            }
        }, ensure_ascii=False) + "\n"

    except Exception as e:
        tb = traceback.format_exc()
        yield json.dumps({
            "type": "error",
            "message": f"评分失败（{current_stage}）：{str(e)}",
            "detail": tb,
            "stage": current_stage,
            "available_models": [{"id": m.get("id"), "name": m.get("name")} for m in get_all_enabled_model_configs()]
        }, ensure_ascii=False) + "\n"
    finally:
        if upload_temp_dir and os.path.exists(upload_temp_dir):
            import shutil
            try:
                shutil.rmtree(upload_temp_dir)
            except:
                pass


def _score_single_bid(bid_file, criteria, model_config, model_name, supports_vision):
    from backend.bid_parser import parse_bid_file, match_section_to_criterion, find_images_for_criterion, compress_image
    from backend.markdown_converter import detect_file_type

    bid_filename = bid_file.get("filename", "")
    bid_path = bid_file.get("path", "")
    
    all_scores = []
    total_score = 0
    llm_call_count = 0
    image_verifications = 0

    bid_type = detect_file_type(bid_path)
    parse_result = parse_bid_file(bid_path, bid_type)

    if not parse_result["success"]:
        return {
            "success": False,
            "filename": bid_filename,
            "error": parse_result.get("error", "未知错误"),
            "scores": [],
        }

    sections = parse_result["sections"]
    raw_text = parse_result["raw_text"]
    images = parse_result["images"]

    for criterion in criteria:
        if criterion["category"] == "price":
            score_result = _score_price(criterion, raw_text, model_name, model_config)
        elif criterion["category"] == "commercial":
            score_result = _score_commercial(criterion, sections, images, raw_text, supports_vision, model_name, model_config)
            if criterion.get("requires_image", False) and score_result.get("image_evidence"):
                image_verifications += 1
        else:
            score_result = _score_technical(criterion, sections, raw_text, model_name, model_config)

        llm_call_count += score_result.get("llm_calls", 1)
        all_scores.append(score_result)
        total_score += score_result["score"]

    strengths, weaknesses = _generate_strengths_weaknesses(all_scores)

    return {
        "success": True,
        "filename": bid_filename,
        "total_score": round(total_score, 2),
        "total_max": sum(c["max_score"] for c in criteria),
        "scores": all_scores,
        "strengths": strengths,
        "weaknesses": weaknesses,
        "llm_calls": llm_call_count,
        "image_verifications": image_verifications,
        "sections_count": len(sections),
        "images_count": len(images),
    }


def _score_price(criterion, bid_text, model_name, model_config):
    from backend.llm_client import call_llm
    from backend.ai_extractor import _call_llm_with_fallback

    system_prompt = """你是一位专业的报价分析专家。请从投标文件中提取投标报价并进行评分。
只返回 JSON 对象本身，不要包含 ``` 代码块或任何解释文字。

输出格式：
{
  "score": 得分,
  "max_score": 满分,
  "reason": "打分依据",
  "evidence": "投标文件中的具体引用",
  "confidence": 0.0-1.0,
  "bid_price": 投标报价数字（如 1250000.0）
}

要求：
1. 根据投标报价的合理性、完整性打分
2. 如果投标文件中没有明确的报价信息，score 为 0
3. evidence 必须引用投标文件中的具体原文
"""

    user_content = f"评分标准：{criterion['name']}（满分 {criterion['max_score']} 分）\n\n==== 投标文件内容 ====\n{_truncate(bid_text, 80000)}"

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_content},
    ]

    try:
        if model_config:
            response = call_llm(messages, temperature=0.1, max_tokens=2000, model_config=model_config)
        else:
            response, model_name = _call_llm_with_fallback(messages, temperature=0.1, max_tokens=2000)

        parsed = _parse_json_response(response)

        score = float(parsed.get("score", 0) or 0)
        max_score = float(parsed.get("max_score", criterion["max_score"]) or criterion["max_score"])
        score = min(score, max_score)

        return {
            "id": criterion["id"],
            "name": criterion["name"],
            "max_score": max_score,
            "score": round(score, 2),
            "category": "price",
            "reason": str(parsed.get("reason", "")),
            "evidence": str(parsed.get("evidence", "")),
            "confidence": float(parsed.get("confidence", 0.5) or 0.5),
            "evidence_type": "text",
        }

    except Exception as e:
        return {
            "id": criterion["id"],
            "name": criterion["name"],
            "max_score": criterion["max_score"],
            "score": 0,
            "category": "price",
            "reason": f"评分失败：{str(e)}",
            "evidence": "",
            "confidence": 0.1,
            "evidence_type": "text",
        }


def _score_commercial(criterion, sections, images, raw_text, supports_vision, model_name, model_config):
    from backend.llm_client import call_llm, encode_image_to_base64, _ocr_image_to_text
    from backend.ai_extractor import _call_llm_with_fallback
    from backend.bid_parser import find_images_for_criterion, compress_image, match_section_to_criterion

    if criterion.get("requires_image", False) and images:
        relevant_images = find_images_for_criterion(criterion, images, sections)
        if relevant_images:
            results = []
            for img in relevant_images[:3]:
                compressed_data = compress_image(img["data"], max_size_kb=500)
                img_result = _verify_image_evidence(compressed_data, img["filename"], criterion, supports_vision, model_name, model_config)
                results.append(img_result)

            if results:
                valid_count = sum(1 for r in results if r.get("is_valid", False))
                avg_confidence = sum(r.get("confidence", 0) for r in results) / len(results)

                if valid_count > 0:
                    score_ratio = min(1.0, valid_count / max(1, len(criterion.get("sub_items", [])) or 1))
                    score = round(criterion["max_score"] * score_ratio, 2)
                else:
                    score = 0

                evidence_text = "; ".join(f"{r.get('cert_name', '')}({r.get('expire_date', '')})" for r in results)
                reason = f"验证 {len(results)} 张图片证据，{valid_count} 张有效"

                return {
                    "id": criterion["id"],
                    "name": criterion["name"],
                    "max_score": criterion["max_score"],
                    "score": min(score, criterion["max_score"]),
                    "category": "commercial",
                    "reason": reason,
                    "evidence": evidence_text,
                    "confidence": round(avg_confidence, 2),
                    "evidence_type": "image",
                    "images_verified": results,
                }

    matched_sections = match_section_to_criterion({"title": "", "content": raw_text}, [criterion])
    if matched_sections:
        system_prompt = """你是一位专业的评标专家。请根据评分标准，对照投标文件内容给出评分。
只返回 JSON 对象本身，不要包含 ``` 代码块或任何解释文字。

输出格式：
{
  "score": 得分,
  "max_score": 满分,
  "reason": "打分依据",
  "evidence": "投标文件中的具体引用",
  "confidence": 0.0-1.0,
  "match_details": [{"keyword": "关键词", "found": true或false, "context": "匹配上下文"}]
}

要求：
1. 根据内容匹配程度、完整性、准确性综合打分
2. evidence 必须引用投标文件中的具体原文
3. match_details 列出每个关键词的匹配情况
"""

        user_content = f"评分标准：{criterion['name']}（满分 {criterion['max_score']} 分）\n描述：{criterion.get('description', '')}\n关键词：{','.join(criterion.get('keywords', []))}\n\n==== 投标文件内容 ====\n{_truncate(raw_text, 80000)}"

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content},
        ]

        try:
            if model_config:
                response = call_llm(messages, temperature=0.1, max_tokens=2000, model_config=model_config)
            else:
                response, model_name = _call_llm_with_fallback(messages, temperature=0.1, max_tokens=2000)

            parsed = _parse_json_response(response)

            score = float(parsed.get("score", 0) or 0)
            max_score = float(parsed.get("max_score", criterion["max_score"]) or criterion["max_score"])
            score = min(score, max_score)

            return {
                "id": criterion["id"],
                "name": criterion["name"],
                "max_score": max_score,
                "score": round(score, 2),
                "category": "commercial",
                "reason": str(parsed.get("reason", "")),
                "evidence": str(parsed.get("evidence", "")),
                "confidence": float(parsed.get("confidence", 0.5) or 0.5),
                "evidence_type": "text",
            }

        except Exception as e:
            pass

    return {
        "id": criterion["id"],
        "name": criterion["name"],
        "max_score": criterion["max_score"],
        "score": 0,
        "category": "commercial",
        "reason": "未找到相关证据或内容",
        "evidence": "",
        "confidence": 0.1,
        "evidence_type": "text",
    }


def _verify_image_evidence(image_data, filename, criterion, supports_vision, model_name, model_config):
    from backend.llm_client import call_llm, encode_image_to_base64, _ocr_image_to_text
    from backend.ai_extractor import _call_llm_with_fallback

    system_prompt = """你是一位专业的资质证书验证专家。请识别图片中的证书内容并验证其有效性。
只返回 JSON 对象本身，不要包含 ``` 代码块或任何解释文字。

输出格式：
{
  "cert_name": "证书全称",
  "cert_type": "证书类型",
  "issuer": "颁发机构",
  "expire_date": "到期日期(YYYY-MM-DD)",
  "is_valid": true或false,
  "confidence": 0.0-1.0,
  "extracted_text": "图片中提取的关键文字"
}

要求：
1. 如果未明确标注有效期，则认为在有效期内（is_valid=true）
2. 日期格式统一为 YYYY-MM-DD
3. confidence 表示识别的置信度
4. extracted_text 提取图片中的关键文字信息
"""

    if supports_vision:
        image_url = encode_image_to_base64(image_data)
        messages = [
            {"role": "system", "content": system_prompt},
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": f"评分标准：{criterion['name']}\n描述：{criterion.get('description', '')}\n请验证这张图片是否符合评分要求"},
                    {"type": "image_url", "image_url": {"url": image_url}},
                ],
            },
        ]
    else:
        ocr_text = _ocr_image_to_text(image_data)
        if not ocr_text:
            return {
                "filename": filename,
                "cert_name": "",
                "cert_type": "",
                "issuer": "",
                "expire_date": "",
                "is_valid": False,
                "confidence": 0.0,
                "extracted_text": "OCR 识别失败，当前模型不支持视觉识别",
            }

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"评分标准：{criterion['name']}\n描述：{criterion.get('description', '')}\n\n请从以下 OCR 提取的内容中验证证书信息：\n\n{ocr_text[:6000]}"},
        ]

    try:
        if model_config:
            response = call_llm(messages, temperature=0.0, max_tokens=1000, model_config=model_config)
        else:
            response, model_name = _call_llm_with_fallback(messages, temperature=0.0, max_tokens=1000)

        parsed = _parse_json_response(response)

        return {
            "filename": filename,
            "cert_name": str(parsed.get("cert_name", "")),
            "cert_type": str(parsed.get("cert_type", "")),
            "issuer": str(parsed.get("issuer", "")),
            "expire_date": str(parsed.get("expire_date", "")),
            "is_valid": bool(parsed.get("is_valid", False)),
            "confidence": float(parsed.get("confidence", 0.5) or 0.5),
            "extracted_text": str(parsed.get("extracted_text", "")),
        }

    except Exception as e:
        return {
            "filename": filename,
            "cert_name": "",
            "cert_type": "",
            "issuer": "",
            "expire_date": "",
            "is_valid": False,
            "confidence": 0.1,
            "extracted_text": f"验证失败：{str(e)}",
        }


def _score_technical(criterion, sections, raw_text, model_name, model_config):
    from backend.llm_client import call_llm
    from backend.ai_extractor import _call_llm_with_fallback

    tech_content = ""
    for section in sections:
        section_text = section.get("content", "")
        for kw in criterion.get("keywords", []):
            if kw in section_text or kw in section.get("title", ""):
                tech_content += f"【{section.get('title', '')}】\n{section_text}\n\n"
                break

    if not tech_content.strip():
        tech_content = raw_text

    system_prompt = """你是一位专业的评标专家。请根据评分标准，对照投标文件技术内容给出评分。
只返回 JSON 对象本身，不要包含 ``` 代码块或任何解释文字。

输出格式：
{
  "score": 得分,
  "max_score": 满分,
  "reason": "打分依据",
  "evidence": "投标文件中的具体引用",
  "confidence": 0.0-1.0,
  "strengths": ["优势点"],
  "weaknesses": ["不足点"],
  "improvement_suggestions": ["改进建议"]
}

要求：
1. 根据内容匹配程度、完整性、技术可行性综合打分
2. evidence 必须引用投标文件中的具体原文
3. confidence 表示对该项打分的把握程度
4. 如果投标文件在某项评分标准下完全没有相关内容，score 为 0
"""

    sub_items_text = ""
    if criterion.get("sub_items"):
        sub_items_text = "\n子项要求：" + "; ".join(f"{si['name']}（{si['max_score']}分）" for si in criterion["sub_items"])

    user_content = f"评分标准：{criterion['name']}（满分 {criterion['max_score']} 分）{sub_items_text}\n描述：{criterion.get('description', '')}\n关键词：{','.join(criterion.get('keywords', []))}\n\n==== 投标文件相关内容 ====\n{_truncate(tech_content, 80000)}"

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_content},
    ]

    try:
        if model_config:
            response = call_llm(messages, temperature=0.1, max_tokens=3000, model_config=model_config)
        else:
            response, model_name = _call_llm_with_fallback(messages, temperature=0.1, max_tokens=3000)

        parsed = _parse_json_response(response)

        score = float(parsed.get("score", 0) or 0)
        max_score = float(parsed.get("max_score", criterion["max_score"]) or criterion["max_score"])
        score = min(score, max_score)

        return {
            "id": criterion["id"],
            "name": criterion["name"],
            "max_score": max_score,
            "score": round(score, 2),
            "category": "technical",
            "reason": str(parsed.get("reason", "")),
            "evidence": str(parsed.get("evidence", "")),
            "confidence": float(parsed.get("confidence", 0.5) or 0.5),
            "evidence_type": "text",
            "strengths": [str(x).strip() for x in (parsed.get("strengths") or []) if str(x).strip()],
            "weaknesses": [str(x).strip() for x in (parsed.get("weaknesses") or []) if str(x).strip()],
            "improvement_suggestions": [str(x).strip() for x in (parsed.get("improvement_suggestions") or []) if str(x).strip()],
        }

    except Exception as e:
        return {
            "id": criterion["id"],
            "name": criterion["name"],
            "max_score": criterion["max_score"],
            "score": 0,
            "category": "technical",
            "reason": f"评分失败：{str(e)}",
            "evidence": "",
            "confidence": 0.1,
            "evidence_type": "text",
            "strengths": [],
            "weaknesses": [],
            "improvement_suggestions": [],
        }


@bid_score_bp.route("/api/bid-score/download", methods=["POST"])
def download_bid_score_report():
    data = request.get_json(silent=True) or {}
    report_format = data.get("format", "docx")
    
    if not data.get("bid_results") and not data.get("scores"):
        return jsonify({"success": False, "message": "报告内容为空"}), 400

    try:
        if report_format == "html":
            html_content = _build_score_report_html(data)
            bio = io.BytesIO(html_content.encode('utf-8'))
            fname = f"模拟评分报告_{datetime.now().strftime('%Y%m%d%H%M%S')}.html"
            mimetype = "text/html"
        else:
            bio = _build_score_report_docx(data)
            fname = f"模拟评分报告_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
            mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    except Exception as e:
        return jsonify({"success": False, "message": "报告生成失败：" + str(e)}), 500

    return send_file(bio, as_attachment=True, download_name=fname, mimetype=mimetype)


def _build_score_report_docx(data):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    normal = doc.styles['Normal']
    normal.font.name = '宋体'
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    title = doc.add_heading('投标文件模拟评分报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run('投标文件：').bold = True
    meta.add_run(data.get('file_name', '—') + '\n')
    meta.add_run('评分时间：').bold = True
    meta.add_run(data.get('checked_at', datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
    for run in meta.runs:
        run.font.size = Pt(10)

    doc.add_paragraph('')

    doc.add_heading('一、评分结果', level=1)

    total_score = data.get('total_score', 0)
    total_max = data.get('total_max', 100)
    score_p = doc.add_paragraph()
    score_run = score_p.add_run(f"总得分：{total_score} / {total_max} 分")
    score_run.bold = True
    score_run.font.size = Pt(14)
    score_run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    doc.add_paragraph('')

    doc.add_heading('二、分项评分详情', level=1)
    scores = data.get("scores") or []

    if scores:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "序号"
        hdr[1].text = "评分项"
        hdr[2].text = "满分"
        hdr[3].text = "得分"
        hdr[4].text = "评分说明"
        for cell in hdr:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        for i, s in enumerate(scores, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = s.get("name", "—")
            row[2].text = str(s.get("max_score", 0))
            row[3].text = str(s.get("score", 0))
            row[4].text = s.get("reason", "")[:100]
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)

        doc.add_paragraph('')

    doc.add_heading('三、评分详情', level=1)
    for s in scores:
        name = s.get("name", "—")
        sc = s.get("score", 0)
        mx = s.get("max_score", 0)

        p = doc.add_paragraph()
        run = p.add_run(f"【{name}】{sc} / {mx} 分")
        run.bold = True
        run.font.size = Pt(12)

        reason = s.get("reason", "")
        if reason:
            rp = doc.add_paragraph()
            rp.add_run("评分依据：").bold = True
            rp.add_run(reason)

        evidence = s.get("evidence", "")
        if evidence:
            ep = doc.add_paragraph()
            ep.add_run("引用内容：").bold = True
            er = ep.add_run(evidence[:200])
            er.font.size = Pt(10)
            er.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        images_verified = s.get("images_verified")
        if images_verified:
            ip = doc.add_paragraph()
            ip.add_run("图片验证结果：").bold = True
            for img in images_verified:
                img_p = doc.add_paragraph()
                status = "✓ 有效" if img.get("is_valid") else "✗ 无效"
                img_p.add_run(f"  • {img.get('filename', '')}: {img.get('cert_name', '')} {status}")
                img_p.font.size = Pt(10)

        strengths = s.get("strengths") or []
        if strengths:
            sp = doc.add_paragraph()
            sp.add_run("优势：").bold = True
            sp.add_run("；".join(strengths))

        weaknesses = s.get("weaknesses") or []
        if weaknesses:
            wp = doc.add_paragraph()
            wp.add_run("不足：").bold = True
            wr = wp.add_run("；".join(weaknesses))
            wr.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

        suggestions = s.get("improvement_suggestions") or []
        if suggestions:
            sp = doc.add_paragraph()
            sp.add_run("改进建议：").bold = True
            sr = sp.add_run("；".join(suggestions))
            sr.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)

        doc.add_paragraph('')

    audit = data.get("audit") or {}
    if audit:
        doc.add_heading('四、审计信息', level=1)
        audit_items = [
            ("评分模型", audit.get("scoring_model", "—")),
            ("评分开始时间", audit.get("started_at", "—")),
            ("评分完成时间", audit.get("finished_at", "—")),
            ("LLM 调用次数", str(audit.get("total_llm_calls", "—"))),
            ("评分项数量", str(audit.get("criteria_count", "—"))),
        ]
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for label, value in audit_items:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)
            for run in row[0].paragraphs[0].runs:
                run.bold = True
        doc.add_paragraph('')

    doc.add_heading('五、重要声明', level=1)
    p = doc.add_paragraph()
    rn = p.add_run('⚠ 模拟评分仅供参考，实际评分以评审专家为准。')
    rn.font.size = Pt(11)
    rn.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    rn.bold = True

    footer = doc.add_paragraph()
    rn = footer.add_run('本报告由 BidTool 投标工具 自动生成。')
    rn.font.size = Pt(9)
    rn.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def _build_score_report_html(data):
    bid_results = data.get("bid_results") or []
    criteria = data.get("criteria") or []
    audit = data.get("audit") or {}
    timestamp = data.get("timestamp", "")

    html_template = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>投标文件模拟评分报告</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: "Microsoft YaHei", "PingFang SC", sans-serif; color: #333; background: #f5f5f5; padding: 40px; }}
        .report-container {{ max-width: 1200px; margin: 0 auto; background: white; padding: 40px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        .report-title {{ text-align: center; font-size: 24px; font-weight: bold; color: #1a73e8; margin-bottom: 30px; border-bottom: 2px solid #1a73e8; padding-bottom: 15px; }}
        .report-meta {{ text-align: center; margin-bottom: 30px; font-size: 14px; color: #666; }}
        .report-meta span {{ margin: 0 15px; }}
        
        .section-title {{ font-size: 18px; font-weight: bold; color: #1a73e8; margin: 30px 0 20px; padding-left: 10px; border-left: 4px solid #1a73e8; }}
        
        .comparison-table {{ width: 100%; border-collapse: collapse; margin-bottom: 30px; }}
        .comparison-table th, .comparison-table td {{ border: 1px solid #ddd; padding: 12px 15px; text-align: center; font-size: 13px; }}
        .comparison-table th {{ background: #f0f4f8; font-weight: bold; color: #333; }}
        .comparison-table th.criteria-col {{ text-align: left; width: 20%; }}
        .comparison-table th.bidder-col {{ background: #e3f2fd; color: #1565c0; }}
        .comparison-table td.criteria-name {{ text-align: left; font-weight: 500; }}
        .comparison-table td.score-high {{ color: #1565c0; font-weight: bold; }}
        .comparison-table td.score-medium {{ color: #f57c00; font-weight: bold; }}
        .comparison-table td.score-low {{ color: #d32f2f; font-weight: bold; }}
        .comparison-table tr.total-row {{ background: #e3f2fd; font-weight: bold; }}
        
        .bidder-detail {{ margin-bottom: 30px; padding: 20px; border: 1px solid #e0e0e0; border-radius: 8px; }}
        .bidder-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 20px; padding-bottom: 15px; border-bottom: 1px solid #e0e0e0; }}
        .bidder-name {{ font-size: 16px; font-weight: bold; color: #333; }}
        .bidder-total {{ font-size: 20px; font-weight: bold; color: #1565c0; }}
        
        .score-item {{ margin-bottom: 15px; padding: 15px; background: #fafafa; border-radius: 6px; }}
        .score-item-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 8px; }}
        .score-item-name {{ font-weight: 600; font-size: 14px; }}
        .score-item-score {{ font-size: 16px; font-weight: bold; }}
        .score-item-reason {{ font-size: 13px; color: #666; margin-top: 5px; }}
        .score-item-evidence {{ font-size: 12px; color: #999; margin-top: 5px; font-style: italic; }}
        
        .evaluation-tags {{ margin-top: 15px; }}
        .evaluation-tag {{ display: inline-block; padding: 4px 12px; border-radius: 15px; font-size: 12px; margin-right: 8px; margin-bottom: 8px; }}
        .tag-strength {{ background: #e8f5e9; color: #2e7d32; }}
        .tag-weakness {{ background: #ffebee; color: #c62828; }}
        
        .audit-section {{ margin-top: 30px; padding: 20px; background: #f8f9fa; border-radius: 8px; }}
        .audit-table {{ width: 100%; border-collapse: collapse; }}
        .audit-table td {{ padding: 8px 15px; border: 1px solid #ddd; font-size: 13px; }}
        .audit-table td.label {{ font-weight: bold; background: #e9ecef; width: 30%; }}
        
        .disclaimer {{ margin-top: 30px; padding: 15px; background: #fff3e0; border-left: 4px solid #ff9800; font-size: 13px; color: #e65100; }}
        
        .footer {{ text-align: center; margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; font-size: 12px; color: #999; }}
        
        @media print {{
            body {{ background: white; padding: 20px; }}
            .report-container {{ box-shadow: none; }}
        }}
    </style>
</head>
<body>
    <div class="report-container">
        <h1 class="report-title">投标文件模拟评分报告</h1>
        
        <div class="report-meta">
            <span><strong>投标人数量：</strong>{len(bid_results)} 家</span>
            <span><strong>评分时间：</strong>{timestamp or '-'}</span>
        </div>

        <h2 class="section-title">一、投标人评分对比</h2>
        <table class="comparison-table">
            <thead>
                <tr>
                    <th rowspan="2" class="criteria-col">评分项</th>
                    <th rowspan="2">满分</th>
                    {''.join(f'<th colspan="2" class="bidder-col">{br.get("filename", "-")}</th>' for br in bid_results)}
                </tr>
                <tr>
                    {''.join(f'<th>得分</th><th>比例</th>' for _ in bid_results)}
                </tr>
            </thead>
            <tbody>
                {''.join(_html_criteria_row(c, bid_results) for c in criteria)}
                <tr class="total-row">
                    <td class="criteria-name"><strong>总分</strong></td>
                    <td>{sum(c["max_score"] for c in criteria)}</td>
                    {''.join(_html_total_cell(br) for br in bid_results)}
                </tr>
            </tbody>
        </table>

        <h2 class="section-title">二、投标人详细评分</h2>
        {''.join(_html_bidder_detail(br, criteria) for br in bid_results)}

        <div class="audit-section">
            <h2 class="section-title">三、审计信息</h2>
            <table class="audit-table">
                <tr><td class="label">评分模型</td><td>{audit.get("model", "—")}</td></tr>
                <tr><td class="label">LLM 调用次数</td><td>{audit.get("llm_calls", "—")}</td></tr>
                <tr><td class="label">投标人数量</td><td>{audit.get("bid_count", "—")}</td></tr>
                <tr><td class="label">评分项数量</td><td>{audit.get("criteria_count", "—")}</td></tr>
                <tr><td class="label">开始时间</td><td>{audit.get("started_at", "—")}</td></tr>
                <tr><td class="label">完成时间</td><td>{audit.get("finished_at", "—")}</td></tr>
                <tr><td class="label">耗时</td><td>{audit.get("duration", "—")}</td></tr>
            </table>
        </div>

        <div class="disclaimer">
            ⚠ 模拟评分仅供参考，实际评分以评审专家为准。本报告由 BidTool 投标工具自动生成。
        </div>

        <div class="footer">
            本报告由 BidTool 投标工具 自动生成
        </div>
    </div>
</body>
</html>"""
    return html_template


def _html_criteria_row(criterion, bid_results):
    max_score = criterion["max_score"]
    cells = ""
    for br in bid_results:
        scores = br.get("scores") or []
        score = next((s.get("score", 0) for s in scores if s.get("name") == criterion["name"]), 0)
        pct = round(score / max_score * 100, 0) if max_score > 0 else 0
        score_class = "score-high" if pct >= 80 else "score-medium" if pct >= 50 else "score-low"
        cells += f'<td class="{score_class}">{score}</td><td>{pct}%</td>'
    
    return f'<tr><td class="criteria-name">{criterion["name"]}</td><td>{max_score}</td>{cells}</tr>'


def _html_total_cell(bid_result):
    total_score = bid_result.get("total_score", 0)
    total_max = bid_result.get("total_max", 1)
    pct = round(total_score / total_max * 100, 1) if total_max > 0 else 0
    score_class = "score-high" if pct >= 80 else "score-medium" if pct >= 50 else "score-low"
    return f'<td class="{score_class}">{total_score}</td><td>{pct}%</td>'


def _html_bidder_detail(bid_result, criteria):
    filename = bid_result.get("filename", "-")
    total_score = bid_result.get("total_score", 0)
    total_max = bid_result.get("total_max", 0)
    scores = bid_result.get("scores") or []
    strengths = bid_result.get("strengths") or []
    weaknesses = bid_result.get("weaknesses") or []
    
    score_items = ""
    for score in scores:
        max_score = score.get("max_score", 0)
        pct = round(score.get("score", 0) / max_score * 100, 0) if max_score > 0 else 0
        score_class = "score-high" if pct >= 80 else "score-medium" if pct >= 50 else "score-low"
        score_items += f"""
        <div class="score-item">
            <div class="score-item-header">
                <span class="score-item-name">{score.get("name", "-")}</span>
                <span class="score-item-score {score_class}">{score.get("score", 0)} / {max_score}</span>
            </div>
            {f'<div class="score-item-reason"><strong>打分依据：</strong>{score.get("reason", "")}</div>' if score.get("reason") else ''}
            {f'<div class="score-item-evidence"><strong>引用内容：</strong>{score.get("evidence", "")}</div>' if score.get("evidence") else ''}
        </div>
        """
    
    strength_tags = ''.join(f'<span class="evaluation-tag tag-strength">{s}</span>' for s in strengths)
    weakness_tags = ''.join(f'<span class="evaluation-tag tag-weakness">{w}</span>' for w in weaknesses)
    
    return f"""
    <div class="bidder-detail">
        <div class="bidder-header">
            <div class="bidder-name">{filename}</div>
            <div class="bidder-total">{total_score} / {total_max}</div>
        </div>
        {score_items}
        <div class="evaluation-tags">
            {f'<div><strong>优势：</strong>{strength_tags}</div>' if strengths else ''}
            {f'<div><strong>不足：</strong>{weakness_tags}</div>' if weaknesses else ''}
        </div>
    </div>"""
