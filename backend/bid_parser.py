import os
import re
import io
from typing import List, Dict, Optional, Tuple
from PIL import Image


def parse_bid_file(file_path: str, file_type: str) -> Dict:
    """
    解析投标文件，拆分章节并提取图片

    :param file_path: 文件路径
    :param file_type: 文件类型 (DOCX/PDF/IMAGE/TEXT)
    :return: {
        "sections": [{title, content, images: [{data, filename, page}]}],
        "raw_text": "完整文本",
        "images": [{data, filename, section_title}],
        "success": bool,
        "error": str or None
    }
    """
    result = {
        "sections": [],
        "raw_text": "",
        "images": [],
        "success": False,
        "error": None
    }

    if not os.path.exists(file_path):
        result["error"] = f"文件不存在: {file_path}"
        return result

    try:
        if file_type == "DOCX":
            docx_data = _parse_docx(file_path)
            result["sections"] = docx_data["sections"]
            result["raw_text"] = docx_data["raw_text"]
            result["images"] = docx_data["images"]
        elif file_type == "PDF":
            pdf_data = _parse_pdf(file_path)
            result["sections"] = pdf_data["sections"]
            result["raw_text"] = pdf_data["raw_text"]
            result["images"] = pdf_data["images"]
        elif file_type == "IMAGE":
            image_data = _parse_image(file_path)
            result["sections"] = image_data["sections"]
            result["raw_text"] = image_data["raw_text"]
            result["images"] = image_data["images"]
        elif file_type == "TEXT":
            text_data = _parse_text(file_path)
            result["sections"] = text_data["sections"]
            result["raw_text"] = text_data["raw_text"]
            result["images"] = []
        else:
            result["error"] = f"不支持的文件类型: {file_type}"
            return result

        result["success"] = True

    except Exception as e:
        result["error"] = f"解析失败: {str(e)}"

    return result


def _parse_docx(file_path: str) -> Dict:
    """解析 DOCX 文件，提取章节和图片"""
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        doc = Document(file_path)
        paragraphs = []
        images = []
        image_index = 0

        for rel in doc.part.rels.values():
            if rel.reltype == RT.IMAGE:
                image_data = rel.target_part.blob
                ext = os.path.splitext(rel.target_ref)[1].lower() or '.png'
                filename = f"image_{image_index}{ext}"
                image_index += 1

                images.append({
                    "data": image_data,
                    "filename": filename,
                    "section_title": "未知章节"
                })

        current_section = {"title": "正文", "content": [], "images": []}
        sections = [current_section]

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name.lower()
            if 'heading' in style or '标题' in style:
                level = 1
                if '2' in style or '二' in style:
                    level = 2
                elif '3' in style or '三' in style:
                    level = 3

                current_section = {"title": text, "content": [], "images": [], "level": level}
                sections.append(current_section)
            else:
                current_section["content"].append(text)

        for table in doc.tables:
            table_lines = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_lines.append(' | '.join(row_data))
            if table_lines:
                current_section["content"].append('\n'.join(table_lines))

        processed_sections = []
        for sec in sections:
            if sec["content"]:
                processed_sections.append({
                    "title": sec["title"],
                    "content": '\n'.join(sec["content"]),
                    "images": [],
                    "level": sec.get("level", 1)
                })

        raw_text = '\n\n'.join(sec["content"] for sec in processed_sections)

        return {
            "sections": processed_sections,
            "raw_text": raw_text,
            "images": images
        }

    except Exception as e:
        return {
            "sections": [],
            "raw_text": f"DOCX 解析失败: {str(e)}",
            "images": []
        }


def _parse_pdf(file_path: str) -> Dict:
    """解析 PDF 文件，提取章节和图片"""
    try:
        import pdfplumber
        import PyPDF2

        sections = []
        images = []
        all_text = []
        image_index = 0

        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    all_text.append(text)

                tables = page.extract_tables()
                for table in tables:
                    table_lines = []
                    for row in table:
                        row_data = [str(cell).strip() if cell else '' for cell in row]
                        table_lines.append(' | '.join(row_data))
                    if table_lines:
                        all_text.append('\n'.join(table_lines))

                try:
                    for img in page.images:
                        raw_data = img['stream'].get_data()
                        ext = img.get('ext', 'png')
                        filename = f"page_{page_num + 1}_image_{image_index}.{ext}"
                        image_index += 1

                        if isinstance(raw_data, bytes):
                            image_data = raw_data
                        elif hasattr(raw_data, 'read'):
                            image_data = raw_data.read()
                        else:
                            image_data = bytes(raw_data)

                        images.append({
                            "data": image_data,
                            "filename": filename,
                            "section_title": f"第{page_num + 1}页",
                            "page": page_num + 1
                        })
                except Exception:
                    pass

        raw_text = '\n\n'.join(all_text)

        lines = raw_text.split('\n')
        current_section = {"title": "正文", "content": [], "images": []}
        sections.append(current_section)

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            heading_match = re.match(r'^\d+[\.、]\s+(.+)$', stripped)
            if heading_match and len(stripped) < 50:
                current_section = {"title": heading_match.group(1), "content": [], "images": []}
                sections.append(current_section)
            elif re.match(r'^[（\(]?[一二三四五六七八九十]+[）\)]?\s+(.+)$', stripped) and len(stripped) < 50:
                current_section = {"title": stripped, "content": [], "images": []}
                sections.append(current_section)
            else:
                current_section["content"].append(line)

        processed_sections = []
        for sec in sections:
            if sec["content"]:
                processed_sections.append({
                    "title": sec["title"],
                    "content": '\n'.join(sec["content"]),
                    "images": []
                })

        return {
            "sections": processed_sections,
            "raw_text": raw_text,
            "images": images
        }

    except Exception as e:
        return {
            "sections": [],
            "raw_text": f"PDF 解析失败: {str(e)}",
            "images": []
        }


def _parse_image(file_path: str) -> Dict:
    """解析图片文件（OCR）"""
    try:
        import pytesseract
        from PIL import Image

        img = Image.open(file_path)
        text = pytesseract.image_to_string(img, lang="chi_sim+eng")

        with open(file_path, "rb") as f:
            image_data = f.read()

        return {
            "sections": [{
                "title": os.path.basename(file_path),
                "content": text,
                "images": []
            }],
            "raw_text": text,
            "images": [{
                "data": image_data,
                "filename": os.path.basename(file_path),
                "section_title": os.path.basename(file_path)
            }]
        }

    except Exception as e:
        with open(file_path, "rb") as f:
            image_data = f.read()

        return {
            "sections": [{
                "title": os.path.basename(file_path),
                "content": f"图片 OCR 失败: {str(e)}",
                "images": []
            }],
            "raw_text": f"图片 OCR 失败: {str(e)}",
            "images": [{
                "data": image_data,
                "filename": os.path.basename(file_path),
                "section_title": os.path.basename(file_path)
            }]
        }


def _parse_text(file_path: str) -> Dict:
    """解析纯文本文件"""
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split('\n')
    current_section = {"title": "正文", "content": [], "images": []}
    sections = [current_section]

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        heading_match = re.match(r'^\d+[\.、]\s+(.+)$', stripped)
        if heading_match and len(stripped) < 50:
            current_section = {"title": heading_match.group(1), "content": [], "images": []}
            sections.append(current_section)
        else:
            current_section["content"].append(line)

    processed_sections = []
    for sec in sections:
        if sec["content"]:
            processed_sections.append({
                "title": sec["title"],
                "content": '\n'.join(sec["content"]),
                "images": []
            })

    return {
        "sections": processed_sections,
        "raw_text": content,
        "images": []
    }


def match_section_to_criterion(section: Dict, criteria: List[Dict]) -> List[Dict]:
    """
    根据关键词匹配章节到评分项

    :param section: 章节 {title, content}
    :param criteria: 评分标准列表
    :return: 匹配的评分项列表
    """
    matched = []
    section_text = (section.get("title", "") + " " + section.get("content", "")).lower()

    for criterion in criteria:
        keywords = criterion.get("keywords", [])
        if not keywords:
            continue

        matched_count = 0
        for kw in keywords:
            if kw.lower() in section_text:
                matched_count += 1

        if matched_count > 0:
            matched.append({
                "criterion": criterion,
                "match_count": matched_count,
                "match_score": matched_count / len(keywords)
            })

    return sorted(matched, key=lambda x: x["match_score"], reverse=True)


def find_images_for_criterion(criterion: Dict, images: List[Dict], sections: List[Dict]) -> List[Dict]:
    """
    为评分项查找相关图片

    :param criterion: 评分项
    :param images: 图片列表
    :param sections: 章节列表
    :return: 相关图片列表
    """
    result = []
    keywords = criterion.get("keywords", [])
    criterion_text = criterion.get("name", "") + " " + criterion.get("description", "")

    for image in images:
        section_title = image.get("section_title", "")
        if any(kw in section_title for kw in keywords) or \
           any(kw in criterion_text for kw in section_title.split()):
            result.append(image)

    return result


def extract_images_from_docx(docx_path: str) -> List[Dict]:
    """从 DOCX 中提取所有图片"""
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        doc = Document(docx_path)
        images = []
        image_index = 0

        for rel in doc.part.rels.values():
            if rel.reltype == RT.IMAGE:
                image_data = rel.target_part.blob
                ext = os.path.splitext(rel.target_ref)[1].lower() or '.png'
                filename = f"image_{image_index}{ext}"
                image_index += 1

                images.append({
                    "data": image_data,
                    "filename": filename,
                    "section_title": "未知章节"
                })

        return images
    except Exception:
        return []


def extract_images_from_pdf(pdf_path: str) -> List[Dict]:
    """从 PDF 中提取所有图片"""
    try:
        import pdfplumber

        images = []
        image_index = 0

        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    for img in page.images:
                        raw_data = img['stream'].get_data()
                        ext = img.get('ext', 'png')
                        filename = f"page_{page_num + 1}_image_{image_index}.{ext}"
                        image_index += 1

                        if isinstance(raw_data, bytes):
                            image_data = raw_data
                        elif hasattr(raw_data, 'read'):
                            image_data = raw_data.read()
                        else:
                            image_data = bytes(raw_data)

                        images.append({
                            "data": image_data,
                            "filename": filename,
                            "section_title": f"第{page_num + 1}页",
                            "page": page_num + 1
                        })
                except Exception:
                    pass

        return images
    except Exception:
        return []


def compress_image(image_data: bytes, max_size_kb: int = 500) -> bytes:
    """
    压缩图片数据

    :param image_data: 原始图片数据
    :param max_size_kb: 最大大小（KB）
    :return: 压缩后的图片数据
    """
    try:
        img = Image.open(io.BytesIO(image_data))

        target_size = max_size_kb * 1024
        if len(image_data) <= target_size:
            return image_data

        quality = 95
        while len(image_data) > target_size and quality > 10:
            buffer = io.BytesIO()
            img.save(buffer, format='JPEG', quality=quality)
            image_data = buffer.getvalue()
            quality -= 10

        if len(image_data) > target_size:
            max_dimension = 1024
            width, height = img.size
            if width > max_dimension or height > max_dimension:
                if width > height:
                    new_width = max_dimension
                    new_height = int(height * max_dimension / width)
                else:
                    new_height = max_dimension
                    new_width = int(width * max_dimension / height)
                img = img.resize((new_width, new_height), Image.LANCZOS)

                buffer = io.BytesIO()
                img.save(buffer, format='JPEG', quality=80)
                image_data = buffer.getvalue()

        return image_data

    except Exception:
        return image_data
